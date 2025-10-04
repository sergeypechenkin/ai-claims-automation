#extract_text.py v2
#import base64
from ast import If
import logging
import os
#import re
import shutil
#from xmlrpc import client
from docx import Document
import docx2txt
import pdfplumber
from pdf2image import convert_from_path
from PIL import Image
from io import BytesIO
import contextlib
import tempfile
from urllib.parse import urlparse
import requests
import zipfile
from azure.ai.vision.imageanalysis import ImageAnalysisClient
#from azure.ai.vision.imageanalysis.models import VisualFeatures
from azure.core.credentials import AzureKeyCredential
#from azure.core.exceptions import HttpResponseError
from azure.identity import DefaultAzureCredential
import json
from azure.identity import DefaultAzureCredential, get_bearer_token_provider
#from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.core.credentials import AzureKeyCredential
#from azure.ai.documentintelligence.models import AnalyzeDocumentRequest
from openai import AzureOpenAI 
from datetime import datetime, timedelta
from azure.storage.blob import (BlobServiceClient, generate_blob_sas, BlobSasPermissions)
import time
import uuid


def get_gpt5_client():
    gpt5_endpoint = os.getenv("GPT5_ENDPOINT", "").strip()
    gpt5_deployment = os.getenv("GPT5_DEPLOYMENT", "").strip()
    gpt5_model_name = os.getenv("GPT5_MODEL", "").strip()
    gpt5_key = os.getenv("GPT5_KEY", "").strip()
    gpt5_api_version = "2024-12-01-preview"

    if not gpt5_endpoint:
        try:
            with open("local.settings.json", "r") as fh:
                settings = json.load(fh)["Values"]
            gpt5_endpoint = gpt5_endpoint or settings.get("GPT5_ENDPOINT", "")
            gpt5_deployment = gpt5_deployment or settings.get("GPT5_DEPLOYMENT", "")
            gpt5_model_name = gpt5_model_name or settings.get("GPT5_MODEL", "")
            gpt5_key = gpt5_key or settings.get("GPT5_KEY", "")
        except FileNotFoundError:
            pass

    if not gpt5_endpoint:
        raise RuntimeError("GPT5_ENDPOINT is missing")

    if gpt5_key:  
        # 🔑 Locally via API Key
        client = AzureOpenAI(
            api_key=gpt5_key,
            api_version=gpt5_api_version,
            azure_endpoint=gpt5_endpoint
        )
    else:  
        # 🔐 In the cloud via Managed Identity
        credential = DefaultAzureCredential()
        token_provider = get_bearer_token_provider(
            credential, "https://cognitiveservices.azure.com/.default"
        )
        client = AzureOpenAI(
            api_version=gpt5_api_version,
            azure_endpoint=gpt5_endpoint,
            azure_ad_token_provider=token_provider
        )

    return {
        "client": client,
        "deployment": gpt5_deployment,
        "model_name": gpt5_model_name
    }
def get_ai_services_client():
    
    endpoint = os.getenv("AI_SERVICES_ENDPOINT", "").strip()
    key = os.getenv("AI_SERVICES_KEY", "").strip()

    if not endpoint:
        try:
            with open("local.settings.json", "r") as fh:
                settings = json.load(fh)["Values"]
            endpoint = endpoint or settings.get("AI_SERVICES_ENDPOINT", "")
            key = key or settings.get("AI_SERVICES_KEY", "")
        except FileNotFoundError:
            pass

    if not endpoint:
        raise RuntimeError("AI_SERVICES_ENDPOINT is missing")

    if key:  
        credential = AzureKeyCredential(key)
    else:  
        credential = DefaultAzureCredential()

    client = ImageAnalysisClient(endpoint=endpoint, credential=credential)
    return client

def to_blob_sas_url(file_input):
    """
    Convert a blob URI or local file path to a SAS URL with READ permission for 1 hour.
    Args:
        file_input: Either a blob URI like "/emailattachments/filename.docx" or a local file path
    Returns:
        str: SAS URL with READ permission valid for 1 hour
    """
    # Get environment variables
    storage_endpoint = os.getenv("STORAGE_ACCOUNT_BLOB_ENDPOINT", "").strip()
    storage_key = os.getenv("STORAGE_ACCOUNT_KEY", "").strip()
    container_name = os.getenv("EMAIL_ATTACHMENTS_CONTAINER", "emailattachments").strip()
    
    # Fallback to local.settings.json if running locally
    if not storage_endpoint:
        try:
            with open("local.settings.json", "r") as fh:
                settings = json.load(fh)["Values"]
            storage_endpoint = storage_endpoint or settings.get("STORAGE_ACCOUNT_BLOB_ENDPOINT", "")
            storage_key = storage_key or settings.get("STORAGE_ACCOUNT_KEY", "")
            container_name = container_name or settings.get("EMAIL_ATTACHMENTS_CONTAINER", "emailattachments")
        except FileNotFoundError:
            pass
    
    if not storage_endpoint:
        raise RuntimeError("STORAGE_ACCOUNT_BLOB_ENDPOINT is required")
    
    # Use key authentication if available (local), otherwise use managed identity (App Service)
    if storage_key:
        credential = storage_key
    else:
        credential = DefaultAzureCredential()
    
    # Create BlobServiceClient
    blob_service_client = BlobServiceClient(
        account_url=storage_endpoint,
        credential=credential
    )
    
    # Determine if input is a blob URI or local file
    blob_name = None
    
    if file_input.startswith("/"):
        # It's a blob URI like "/emailattachments/filename.docx"
        parts = file_input.lstrip("/").split("/", 1)
        if len(parts) == 2:
            container_name = parts[0]
            blob_name = parts[1]
        else:
            blob_name = parts[0]
    elif os.path.exists(file_input):
        # It's a local file - upload it
        filename = os.path.basename(file_input)
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        blob_name = f"{timestamp}_{filename}"
        
        blob_client = blob_service_client.get_blob_client(
            container=container_name,
            blob=blob_name
        )
        
        with open(file_input, "rb") as data:
            blob_client.upload_blob(data, overwrite=True)
    else:
        # Assume it's already a blob name without leading slash
        blob_name = file_input
    
    # Generate SAS token with READ permission for 1 hour
    # SAS generation requires account key, so we need to handle both cases
    if storage_key:
        sas_token = generate_blob_sas(
            account_name=storage_endpoint.split("//")[1].split(".")[0],
            container_name=container_name,
            blob_name=blob_name,
            account_key=storage_key,
            permission=BlobSasPermissions(read=True),
            expiry=datetime.utcnow() + timedelta(hours=1)
        )
        # Ensure no double slashes in URL
        base_url = storage_endpoint.rstrip("/")
        sas_url = f"{base_url}/{container_name}/{blob_name}?{sas_token}"
    else:
        # When using managed identity, return the blob URL with user delegation SAS
        blob_client = blob_service_client.get_blob_client(
            container=container_name,
            blob=blob_name
        )
        
        # Get user delegation key for generating SAS with managed identity
        delegation_key = blob_service_client.get_user_delegation_key(
            key_start_time=datetime.utcnow(),
            key_expiry_time=datetime.utcnow() + timedelta(hours=1)
        )
        
        sas_token = generate_blob_sas(
            account_name=storage_endpoint.split("//")[1].split(".")[0],
            container_name=container_name,
            blob_name=blob_name,
            user_delegation_key=delegation_key,
            permission=BlobSasPermissions(read=True),
            expiry=datetime.utcnow() + timedelta(hours=1)
        )
        # Ensure no double slashes in URL
        base_url = storage_endpoint.rstrip("/")
        sas_url = f"{base_url}/{container_name}/{blob_name}?{sas_token}"
    
    return sas_url

def get_temp_dir():
    """
    Get appropriate temporary directory based on running environment.
    Returns:
        str: Path to temporary directory
    """
    # Check if running in Azure App Service
    if os.getenv("WEBSITE_INSTANCE_ID"):
        # Use Azure App Service temp folder
        temp_dir = os.getenv("TEMP", "/tmp")
    else:
        # Use system temp directory for local development
        temp_dir = tempfile.gettempdir()
    
    return temp_dir

def download_document_to_temp(blob_uri):
    """
    Download a document from blob storage to temporary local storage.
    Uses managed identity in cloud or key auth locally.
    Args:
        blob_uri: Blob URI (e.g., "/emailattachments/filename.pdf")
    Returns:
        str: Path to downloaded temporary file
    """
    try:
        # Get environment variables
        storage_endpoint = os.getenv("STORAGE_ACCOUNT_BLOB_ENDPOINT", "").strip()
        storage_key = os.getenv("STORAGE_ACCOUNT_KEY", "").strip()
        container_name = os.getenv("EMAIL_ATTACHMENTS_CONTAINER", "emailattachments").strip()
        
        # Fallback to local.settings.json if running locally
        if not storage_endpoint:
            try:
                with open("local.settings.json", "r") as fh:
                    settings = json.load(fh)["Values"]
                storage_endpoint = storage_endpoint or settings.get("STORAGE_ACCOUNT_BLOB_ENDPOINT", "")
                storage_key = storage_key or settings.get("STORAGE_ACCOUNT_KEY", "")
                container_name = container_name or settings.get("EMAIL_ATTACHMENTS_CONTAINER", "emailattachments")
            except FileNotFoundError:
                pass
        
        if not storage_endpoint:
            raise RuntimeError("STORAGE_ACCOUNT_BLOB_ENDPOINT is required")
        
        # Use key authentication if available (local), otherwise use managed identity (cloud)
        if storage_key:
            credential = storage_key
        else:
            credential = DefaultAzureCredential()
        
        # Create BlobServiceClient
        blob_service_client = BlobServiceClient(
            account_url=storage_endpoint,
            credential=credential
        )
        
        # Parse blob URI
        blob_name = None
        if blob_uri.startswith("/"):
            parts = blob_uri.lstrip("/").split("/", 1)
            if len(parts) == 2:
                container_name = parts[0]
                blob_name = parts[1]
            else:
                blob_name = parts[0]
        else:
            blob_name = blob_uri
        
        # Get blob client
        blob_client = blob_service_client.get_blob_client(
            container=container_name,
            blob=blob_name
        )
        
        # Get appropriate temp directory and file extension
        temp_dir = get_temp_dir()
        _, ext = os.path.splitext(blob_name)
        
        # Download to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext, dir=temp_dir) as tmp_file:
            download_stream = blob_client.download_blob()
            tmp_file.write(download_stream.readall())
            temp_path = tmp_file.name
        
        logging.info(f"Downloaded {blob_uri} to {temp_path}")
        return temp_path
    
    except Exception as e:
        logging.error(f"Error downloading document {blob_uri}: {str(e)}")
        raise

def get_image_url(blob_uri):
    """
    Get appropriate URL for image processing.
    Returns SAS URL for local development, blob URL for cloud with managed identity.
    Args:
        blob_uri: Blob URI (e.g., "/emailattachments/filename.jpg")
    Returns:
        str: URL for image processing
    """
    storage_key = os.getenv("STORAGE_ACCOUNT_KEY", "").strip()
    
    # Fallback to local.settings.json
    if not storage_key:
        try:
            with open("local.settings.json", "r") as fh:
                settings = json.load(fh)["Values"]
            storage_key = storage_key or settings.get("STORAGE_ACCOUNT_KEY", "")
        except FileNotFoundError:
            pass
    
    # If running locally with key, generate SAS URL
    if storage_key:
        return to_blob_sas_url(blob_uri)
    
    # In cloud with managed identity, construct blob URL
    storage_endpoint = os.getenv("STORAGE_ACCOUNT_BLOB_ENDPOINT", "").strip()
    container_name = os.getenv("EMAIL_ATTACHMENTS_CONTAINER", "emailattachments").strip()
    
    blob_name = blob_uri.lstrip("/").split("/", 1)[-1]
    return f"{storage_endpoint.rstrip('/')}/{container_name}/{blob_name}"

def extract_docx(file_path):
    """
    Extract text content from DOCX files.
    Args:
        file_path: Local file path to DOCX file
    Returns:
        dict: Extracted content and metadata
    """
    try:
        # Extract text
        doc = Document(file_path)
        text = docx2txt.extract(file_path)
        
        # Extract structured content
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        return {
            "type": "docx",
            "text": text,
            "paragraphs": paragraphs,
            "status": "success"
        }
    except Exception as e:
        logging.error(f"Error extracting DOCX: {str(e)}")
        return {
            "type": "docx",
            "text": "",
            "error": str(e),
            "status": "error"
        }

def extract_doc(file_path):
    """
    Extract text content from DOC files.
    Args:
        file_path: Local file path to DOC file
    Returns:
        dict: Extracted content and metadata
    """
    # DOC files are more complex - would need additional libraries like antiword or conversion
    return {
        "type": "doc",
        "text": "",
        "error": "DOC format requires conversion to DOCX",
        "status": "unsupported"
    }

def extract_pdf(file_path):
    """
    Extract text and images from PDF files.
    Args:
        file_path: Local file path to PDF file
    Returns:
        dict: Extracted content and metadata
    """
    try:
        text_content = []
        page_count = 0
        
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
        
        return {
            "type": "pdf",
            "text": "\n\n".join(text_content),
            "page_count": page_count,
            "status": "success"
        }
    except Exception as e:
        logging.error(f"Error extracting PDF: {str(e)}")
        return {
            "type": "pdf",
            "text": "",
            "error": str(e),
            "status": "error"
        }

def extract_image(image_url):
    """
    Extract text from images using Azure AI Vision.
    Args:
        image_url: SAS URL (local) or blob URL (cloud with managed identity)
    Returns:
        dict: Extracted content and metadata
    """
    try:
        client = get_ai_services_client()
        
        # Analyze image directly from URL
        result = client.analyze_from_url(
            image_url=image_url,
            visual_features=["READ"]
        )
        
        # Extract text
        text_blocks = []
        if result.read and result.read.blocks:
            for block in result.read.blocks:
                for line in block.lines:
                    text_blocks.append(line.text)
        
        return {
            "type": "image",
            "text": "\n".join(text_blocks),
            "status": "success"
        }
    except Exception as e:
        logging.error(f"Error extracting image: {str(e)}")
        return {
            "type": "image",
            "text": "",
            "error": str(e),
            "status": "error"
        }

def extract_attachment_info(attachment_uris):
    """
    Process a list of attachment URIs and extract information based on file type.
    Documents are downloaded to temp storage, images are processed via URL.
    Args:
        attachment_uris: List of blob URIs (e.g., ["/emailattachments/file1.pdf", "/emailattachments/file2.jpg"])
    Returns:
        list: List of dictionaries containing extracted information for each attachment
    """
    results = []
    temp_files_to_cleanup = []
    
    # Document types that need downloading
    document_handlers = {
        '.pdf': extract_pdf,
        '.docx': extract_docx,
        '.doc': extract_doc
    }
    
    # Image types that process via URL
    image_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.tif'}
    
    try:
        for uri in attachment_uris:
            try:
                # Extract file extension from blob URI
                _, ext = os.path.splitext(uri.lower())
                
                # Process documents (download first)
                if ext in document_handlers:
                    logging.info(f"Processing document {ext} file: {uri}")
                    
                    # Download document to temp storage
                    temp_path = download_document_to_temp(uri)
                    temp_files_to_cleanup.append(temp_path)
                    
                    # Extract using local file path
                    handler = document_handlers[ext]
                    result = handler(temp_path)
                    result['uri'] = uri
                    result['extension'] = ext
                    results.append(result)
                
                # Process images (use URL directly)
                elif ext in image_extensions:
                    logging.info(f"Processing image {ext} file: {uri}")
                    
                    # Get appropriate URL (SAS for local, blob URL for cloud)
                    image_url = get_image_url(uri)
                    
                    # Extract using URL
                    result = extract_image(image_url)
                    result['uri'] = uri
                    result['extension'] = ext
                    results.append(result)
                
                else:
                    logging.warning(f"Unsupported file type: {ext} for {uri}")
                    results.append({
                        'uri': uri,
                        'extension': ext,
                        'type': 'unsupported',
                        'text': '',
                        'error': f'Unsupported file type: {ext}',
                        'status': 'unsupported'
                    })
            
            except Exception as e:
                logging.error(f"Error processing attachment {uri}: {str(e)}")
                results.append({
                    'uri': uri,
                    'type': 'error',
                    'text': '',
                    'error': str(e),
                    'status': 'error'
                })
    
    finally:
        # Clean up all temporary files
        for temp_file in temp_files_to_cleanup:
            try:
                if os.path.exists(temp_file):
                    os.unlink(temp_file)
                    logging.info(f"Cleaned up temp file: {temp_file}")
            except Exception as e:
                logging.warning(f"Failed to clean up temp file {temp_file}: {str(e)}")
    
    return results


# Test code - only run when script is executed directly
if __name__ == "__main__":
    # Example usage for testing
    #file_path = "https://staiclaimsauto001.blob.core.windows.net/emailattachments/20251003154607_photos.docx?sp=r&st=2025-10-04T11:03:40Z&se=2025-10-11T19:18:40Z&spr=https&sv=2024-11-04&sr=b&sig=%2FvFXazh2nLE%2BAy0ovcVQTNexNcFRWaP7YznxQZZkt9M%3D"
    #file_path = "/emailattachments/20251004125906_photos.docx"
    #file_path = "C:\\Users\\sergeype\\payload.json"
    #print(to_blob_sas_url(file_path))

    # Test extract_attachment_info with blob URIs
    test_uris = [
    "/emailattachments/20251004172227_Cross RAD1 Referral.doc",
    "/emailattachments/20251004172227_photos.docx"
  ]
    results = extract_attachment_info(test_uris)
    for result in results:
        print(f"\n{result['uri']}:")
        print(f"  Type: {result['type']}")
        print(f"  Status: {result['status']}")
        if result['text']:
            print(f"  Text preview: {result['text'][:100]}...")
            print(f"  Text preview: {result['text'][:100]}...")
