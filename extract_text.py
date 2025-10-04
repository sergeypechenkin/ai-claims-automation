#import base64
import logging
import os
import shutil
import tempfile
from urllib.parse import urlparse
import requests
import zipfile
from docx import Document
import docx2txt
import pdfplumber
from pdf2image import convert_from_path
from PIL import Image
from io import BytesIO
import contextlib
from azure.ai.vision.imageanalysis import ImageAnalysisClient
from azure.core.credentials import AzureKeyCredential
from azure.identity import DefaultAzureCredential, get_bearer_token_provider
from openai import AzureOpenAI 
from datetime import datetime, timedelta
from azure.storage.blob import (BlobServiceClient, generate_blob_sas, BlobSasPermissions)
import time
import uuid
import json



def get_gpt5_client():
    gpt5_endpoint = os.getenv("GPT5_ENDPOINT", "").strip()
    gpt5_deployment = os.getenv("GPT5_DEPLOYMENT", "").strip()
    gpt5_model_name = os.getenv("GPT5_MODEL", "").strip()
    gpt5_key = os.getenv("GPT5_KEY", "").strip()
    gpt5_api_version = "2024-12-01-preview"

    if not gpt5_endpoint:
        try:
            with open("local.settings.json", "r") as fh:
                settings = json.load(fh)["Values"]
            gpt5_endpoint = gpt5_endpoint or settings.get("GPT5_ENDPOINT", "")
            gpt5_deployment = gpt5_deployment or settings.get("GPT5_DEPLOYMENT", "")
            gpt5_model_name = gpt5_model_name or settings.get("GPT5_MODEL", "")
            gpt5_key = gpt5_key or settings.get("GPT5_KEY", "")
        except FileNotFoundError:
            pass

    if not gpt5_endpoint:
        raise RuntimeError("GPT5_ENDPOINT is missing")

    if gpt5_key:  
        # 🔑 Locally via API Key
        client = AzureOpenAI(
            api_key=gpt5_key,
            api_version=gpt5_api_version,
            azure_endpoint=gpt5_endpoint
        )
    else:  
        # 🔐 In the cloud via Managed Identity
        credential = DefaultAzureCredential()
        token_provider = get_bearer_token_provider(
            credential, "https://cognitiveservices.azure.com/.default"
        )
        client = AzureOpenAI(
            api_version=gpt5_api_version,
            azure_endpoint=gpt5_endpoint,
            azure_ad_token_provider=token_provider
        )

    return {
        "client": client,
        "deployment": gpt5_deployment,
        "model_name": gpt5_model_name
    }
def get_ai_services_client():
    endpoint = os.getenv("AI_SERVICES_ENDPOINT", "").strip()
    key = os.getenv("AI_SERVICES_KEY", "").strip()

    if not endpoint:
        try:
            with open("local.settings.json", "r") as fh:
                settings = json.load(fh)["Values"]
            endpoint = endpoint or settings.get("AI_SERVICES_ENDPOINT", "")
            key = key or settings.get("AI_SERVICES_KEY", "")
        except FileNotFoundError:
            pass

    if not endpoint:
        raise RuntimeError("AI_SERVICES_ENDPOINT is missing")

    if key:  
        credential = AzureKeyCredential(key)
    else:  
        credential = DefaultAzureCredential()

    client = ImageAnalysisClient(endpoint=endpoint, credential=credential)
    return client
def _is_remote_path(path: str) -> bool:
    parsed = urlparse(path)
    return parsed.scheme in ("http", "https")
def _resolve_extension(path: str) -> str:
    target = urlparse(path).path if _is_remote_path(path) else path
    return os.path.splitext(target)[1].lower()
def _extract_filename(path: str) -> str:
    parsed = urlparse(path)
    return os.path.basename(parsed.path) if parsed.scheme else os.path.basename(path)
def _download_to_temp(url: str) -> str:
    suffix = os.path.splitext(urlparse(url).path)[1] or ".tmp"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        response = requests.get(url, timeout=30)
        response.raise_for_status()
        tmp.write(response.content)
        return tmp.name

@contextlib.contextmanager
def _ensure_local_file(path: str):
	"""
	Yield a local filesystem path for the given input.
	- If path is an HTTP(S) URL -> download to temp and yield the temp path.
	- If path is a URI-style path that starts with '/' -> convert to SAS URL (ensure_remote_image_url)
	  then download and yield the temp path.
	- Otherwise yield path as-is (assumed to be a local filesystem path).
	"""
	if _is_remote_path(path):
		temp_path = _download_to_temp(path)
		try:
			yield temp_path
		finally:
			if os.path.exists(temp_path):
				os.remove(temp_path)
	elif isinstance(path, str) and path.startswith('/'):
		# treat as storage URI path: convert to https SAS URL then download
		try:
			remote_url = ensure_remote_image_url(path)
		except Exception as exc:
			# bubble up a clear error so callers can handle/log it
			raise RuntimeError(f"Failed to resolve storage URI to remote URL: {exc}") from exc
		temp_path = _download_to_temp(remote_url)
		try:
			yield temp_path
		finally:
			if os.path.exists(temp_path):
				os.remove(temp_path)
	else:
		yield path
def _to_image_bytes(image_source):
    if isinstance(image_source, Image.Image):
        img = image_source.convert("RGB") if image_source.mode not in ("RGB", "RGBA", "L") else image_source
        buffer = BytesIO()
        img.save(buffer, format="PNG")
        data = buffer.getvalue()
        if not data:
            raise ValueError("Failed to serialize image to bytes.")
        return data
    if isinstance(image_source, (bytes, bytearray)):
        return bytes(image_source)
    if isinstance(image_source, str):
        with open(image_source, "rb") as fh:
            return fh.read()
    raise TypeError(f"Unsupported image source type: {type(image_source)!r}")
def _get_storage_account_url():
    # try explicit blob endpoint first, then account name
    url = os.getenv("STORAGE_ACCOUNT_BLOB_ENDPOINT", "").strip()
    
    if not url:
        # local.settings.json fallback
        try:
            with open("local.settings.json", "r") as fh:
                settings = json.load(fh)["Values"]
            url = url or settings.get("STORAGE_ACCOUNT_BLOB_ENDPOINT")

        except Exception:
            pass
    if not url:
        raise RuntimeError("Storage account URL not configured. Set STORAGE_ACCOUNT_BLOB_ENDPOINT.")
    return url

def _upload_file_to_container_and_get_sas(account_url: str, container_name: str, blob_name: str, local_file: str, expiry_hours: int = 1) -> str:
	"""
	Upload a local file to the given container and return a read-only SAS URL.
	(Compact helper used for uploading temp images to a specific container, e.g. "tems".)
	"""
	account_key = os.getenv("STORAGE_ACCOUNT_KEY")
	if not account_key:
		try:
			with open("local.settings.json", "r") as fh:
				settings = json.load(fh)["Values"]
			account_key = account_key or settings.get("STORAGE_ACCOUNT_KEY", "")
		except FileNotFoundError:
			pass

	if account_key:
		blob_service_client = BlobServiceClient(account_url=account_url, credential=account_key)
		user_delegation_key = None
	else:
		credential = DefaultAzureCredential()
		blob_service_client = BlobServiceClient(account_url=account_url, credential=credential)
		user_delegation_key = blob_service_client.get_user_delegation_key(
			key_start_time=datetime.utcnow(),
			key_expiry_time=datetime.utcnow() + timedelta(hours=expiry_hours)
		)

	# Ensure container exists
	container_client = blob_service_client.get_container_client(container_name)
	try:
		container_client.create_container()
	except Exception:
		pass

	blob_client = container_client.get_blob_client(blob_name)
	with open(local_file, "rb") as data:
		blob_client.upload_blob(data, overwrite=True)

	account_name = blob_service_client.account_name or os.getenv("STORAGE_ACCOUNT_NAME")
	if not account_name:
		raise ValueError("Storage account name could not be determined")

	if account_key and blob_service_client.account_name:
		sas_token = generate_blob_sas(
			account_name=account_name,
			container_name=container_name,
			blob_name=blob_name,
			account_key=account_key,
			permission=BlobSasPermissions(read=True),
			expiry=datetime.utcnow() + timedelta(hours=expiry_hours)
		)
	else:
		sas_token = generate_blob_sas(
			account_name=account_name,
			container_name=container_name,
			blob_name=blob_name,
			user_delegation_key=user_delegation_key,
			permission=BlobSasPermissions(read=True),
			expiry=datetime.utcnow() + timedelta(hours=expiry_hours)
		)

	return f"{blob_client.url}?{sas_token}"

def upload_temp_image_and_get_url(local_file_path: str, container: str = "tems") -> str:
	"""
	Upload the given local image to the specified container and return SAS URL.
	"""
	account_url = _get_storage_account_url()
	blob_name = f"{int(time.time())}_{uuid.uuid4().hex}_{os.path.basename(local_file_path)}"
	return _upload_file_to_container_and_get_sas(account_url, container, blob_name, local_file_path)


def ensure_remote_image_url(image_source):
    """
    Return an HTTPS-accessible URL for the given image_source.
    - If image_source is already an http(s) URL -> return as-is.
    - If image_source is a URI path (starts with /) -> construct SAS URL from storage account.
    - If image_source is a local file path -> upload to storage and return SAS URL.
    - If image_source is a PIL.Image -> save to temp file, upload, cleanup, return SAS URL.
    """
    if isinstance(image_source, str):
        if _is_remote_path(image_source):
            return image_source
        
        # Check if it's a URI path (starts with /)
        if image_source.startswith('/'):
            account_url = _get_storage_account_url()
            # Parse container and blob path from URI
            parts = image_source.lstrip('/').split('/', 1)
            if len(parts) == 2:
                container_name, blob_path = parts
            else:
                raise ValueError(f"Invalid URI path format: {image_source}")
            
            # Generate SAS URL for existing blob
            account_key = os.getenv("STORAGE_ACCOUNT_KEY")
            if not account_key:
                try:
                    with open("local.settings.json", "r") as fh:
                        settings = json.load(fh)["Values"]
                    account_key = settings.get("STORAGE_ACCOUNT_KEY", "")
                except FileNotFoundError:
                    pass
            
            if account_key:
                blob_service_client = BlobServiceClient(account_url=account_url, credential=account_key)
                user_delegation_key = None
            else:
                credential = DefaultAzureCredential()
                blob_service_client = BlobServiceClient(account_url=account_url, credential=credential)
                user_delegation_key = blob_service_client.get_user_delegation_key(
                    key_start_time=datetime.utcnow(),
                    key_expiry_time=datetime.utcnow() + timedelta(hours=1)
                )
            
            blob_client = blob_service_client.get_blob_client(container=container_name, blob=blob_path)
            
            account_name = blob_service_client.account_name or os.getenv("STORAGE_ACCOUNT_NAME")
            if not account_name:
                raise ValueError("Storage account name could not be determined")
            
            if account_key:
                sas_token = generate_blob_sas(
                    account_name=account_name,
                    container_name=container_name,
                    blob_name=blob_path,
                    account_key=account_key,
                    permission=BlobSasPermissions(read=True),
                    expiry=datetime.utcnow() + timedelta(hours=1)
                )
            else:
                sas_token = generate_blob_sas(
                    account_name=account_name,
                    container_name=container_name,
                    blob_name=blob_path,
                    user_delegation_key=user_delegation_key,
                    permission=BlobSasPermissions(read=True),
                    expiry=datetime.utcnow() + timedelta(hours=1)
                )
            
            return f"{blob_client.url}?{sas_token}"
        
        # local file path -> upload
        account_url = _get_storage_account_url()
        blob_name = f"tmp_uploads/{int(time.time())}_{os.path.basename(image_source)}"
        sas_url = upload_and_get_sas(account_url, blob_name, image_source)
        return sas_url

    if isinstance(image_source, Image.Image):
        tmp = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as fh:
                tmp = fh.name
                image_source.save(fh, format="PNG")
                fh.flush()
            account_url = _get_storage_account_url()
            blob_name = f"tmp_uploads/{int(time.time())}.png"
            sas_url = upload_and_get_sas(account_url, blob_name, tmp)
            return sas_url
        finally:
            if tmp and os.path.exists(tmp):
                try:
                    os.remove(tmp)
                except Exception:
                    pass

    raise TypeError("ensure_remote_image_url accepts HTTP URL, URI path, local filepath or PIL.Image")
def extract_file_info(file_path, ocr_text_threshold=50):
    ext = _resolve_extension(file_path)
    print(f"\033[93mProcessing file: {file_path} with extension {ext}\033[0m")

    result = {
        "Digital text": "",
        "Images": []   # list with results for images
    }
    text_content = ""
    tables = []
    file_type = "unknown"

    with _ensure_local_file(file_path) as local_path:
        if ext in (".docx", ".doc"):
            is_docx_package = zipfile.is_zipfile(local_path)
            if ext == ".doc" and not is_docx_package:
                result["Summary"] = "Unsupported legacy DOC format."
                logging.warning("Legacy .doc format detected, which is not supported.")
                return result
            if is_docx_package:
                with zipfile.ZipFile(local_path) as archive:
                    entries = {name.lower() for name in archive.namelist()}
                    if "word/document.xml" not in entries:
                        print("Word archive missing document content.")
                        result["Summary"] = "Unsupported Word archive structure."
                        logging.warning("Word archive missing document content.")
                        return result
            doc = None
            if is_docx_package:
                try:
                    doc = Document(local_path)
                except ValueError as exc:
                    print(f"Failed to parse Word document, fallback to docx2txt: {exc}")
                    logging.warning(f"Failed to parse Word document, fallback to docx2txt: {exc}")  
            if doc:
                text_content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        table_data.append([cell.text.strip() for cell in row.cells])
                    tables.append(table_data)
                if tables:
                    print(f"\033[92mExtracted tables from the document: {tables}\033[0m")
                    logging.info(f'Extracted tables from the document: {tables}')
            else:
                print("Docx2txt fallback for text extraction.")
                logging.info("Docx2txt fallback for text extraction.")
                text_content = ""
            img_dir = tempfile.mkdtemp(prefix="docx_images_")
            try:
                extracted_text = ""
                try:
                    extracted_text = docx2txt.process(local_path, img_dir)
                except Exception as exc:
                    print(f"Docx2txt processing failed: {exc}")
                    logging.warning(f"Docx2txt processing failed: {exc}")   
                    result["Summary"] = "Unable to process Word document."
                    return result
                if not doc and extracted_text and not text_content.strip():
                    text_content = extracted_text
                    logging.info(f"Extracted text from Word document: {text_content}")
                img_files = os.listdir(img_dir)
                for img_file in img_files:
                    img_path = os.path.join(img_dir, img_file)
                    try:
                        # upload extracted image to "tems" container and get SAS URL
                        img_url = upload_temp_image_and_get_url(img_path, container="tems")
                        logging.info(f"Uploaded extracted image '{img_path}' -> {img_url}")
                        ocr_text = analyze_image(img_url)
                    except Exception as exc:
                        print(f"Failed to upload/analyze image '{img_path}': {exc}")
                        logging.warning(f"Failed to upload/analyze image '{img_path}': {exc}")
                        ocr_text = ""

                    print(f"Image '{img_file}' analyzed" f" with OCR text: {ocr_text}")
                    logging.info(f"Image '{img_file}' analyzed with OCR text: {ocr_text}")  
                    result["Images"].append({"filename": img_file, "ocr_text": ocr_text})
            except ValueError as exc:
                print(f"Failed to extract Word document content: {exc}")
                result["Summary"] = "Unable to process Word document."
                logging.warning(f"Failed to extract Word document content: {exc}")
                return result
            finally:
                shutil.rmtree(img_dir, ignore_errors=True)
                result["Digital text"] = text_content + "\n" + "\n".join([str(t) for t in tables])
                logging.info(f'Extracted digital text length: {len(text_content)} and tables: {tables}')

        # ----- PDF -----
        elif ext == ".pdf":
            print("PDF file detected")
             # check for text layer
            scanned = True
            with pdfplumber.open(local_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        scanned = False
                        text_content += page_text + "\n"
                        print(f"Extracted text from page {page.page_number}: {page_text}")
                    page_tables = page.extract_tables()
                    for table in page_tables:
                        tables.append(table)
                        print(f"Extracted table from page {page.page_number}: {table}")

            if not scanned:
                file_type = "digital"
                print(f"Document classified as {file_type}.")
            else:
                print("No text layer found, performing OCR.")
                pages = convert_from_path(local_path)
                ocr_chunks = []
                for page in pages:
                    try:
                        page_url = ensure_remote_image_url(page)
                    except Exception as exc:
                        print(f"Failed to ensure remote URL for PDF page image: {exc}")
                        page_text = ""
                    else:
                        page_text = analyze_image(page_url)
                    if page_text:
                         ocr_chunks.append(page_text)
                if ocr_chunks:
                    text_content = "\n".join(ocr_chunks)
                    file_type = "scanned"
        # ----- Images (jpg/png) -----
        elif ext in [".jpg", ".jpeg", ".png", ".tiff"]:
            print("Image file detected")
            logging.info("Image file detected")
            try:
                # Ensure we obtain an HTTPS-accessible URL for the input (handles '/container/blob', local path, or PIL.Image)
                img_url = ensure_remote_image_url(file_path)
                logging.info(f"Image URL for initially image '{file_path}': {img_url}")
            except Exception as exc:
                print(f"Failed to ensure remote URL for image '{file_path}': {exc}")
                logging.warning(f"Failed to ensure remote URL for image '{file_path}': {exc}")
                ocr_text = ""
            else:
                ocr_text = analyze_image(img_url)
            result["Images"].append({
                "filename": _extract_filename(file_path),
                "ocr_text": ocr_text
            })
            print(f"Image OCR text: {ocr_text[:100]}...")
            logging.info(f"Image OCR text: {ocr_text[:100]}...")

        else:
            result["Summary"] = "Unsupported file format."
            logging.warning("Unsupported file format encountered.")
            return result

    # ---------- Анализ текста ----------


    # Convert result dict to string for analysis
    text_for_analysis = json.dumps(result, ensure_ascii=False, indent=2)
    return analyze_text(text_for_analysis)



# def analyze_image(image_source):
#     # https://learn.microsoft.com/en-us/python/api/overview/azure/ai-vision-imageanalysis-readme?view=azure-python#examples
#     # AI Services Vision API works, but performs worse than GPT-5 with images
    
#     client = get_ai_services_client()
#     # Check if image size is less than 100KB
#     image_data = _to_image_bytes(image_source)
#     print(f"Image size: {len(image_data)} bytes")
#     if len(image_data) < 100000:
#         print("Warning: Image is probably a logo or icon, skipping OCR.")
#         return ""
#     try:
#         analysis = client.analyze(image_data=image_data, visual_features=[VisualFeatures.CAPTION, VisualFeatures.READ])
#     except HttpResponseError as exc:
#         print(f"Vision API error: {exc}")
#         return ""

#     lines = []
#     # OCR text
#     if analysis.read and analysis.read.blocks:
#         for block in analysis.read.blocks:
#             for line in block.lines:
#                 lines.append(line.text)
#     ocr_text = ("\n".join(lines))
#     caption_text = analysis.caption.text if analysis.caption else ""
#     print(f"Caption: {caption_text}")
#     print(f"OCR Text: {ocr_text}")

#     if len(ocr_text) > 30:
#         return str("OCR " + ocr_text + "\nCaption: " + caption_text)
#     else:
#         return str("Caption: " + caption_text)

def count_tokens(model_name: str, text: str) -> int:
    """
    Count tokens for the given text.

    Uses the tiktoken library when available (best accuracy). If tiktoken is not
    installed or fails, falls back to a simple heuristic estimate (1 token ≈ 4 chars).
    If a model name is configured (gpt5_model_name or gpt5_deployment) the function
    will attempt to use tiktoken.encoding_for_model(model) to pick the proper encoding.
    """
    try:
        import tiktoken  # optional dependency; import at runtime
        model = None
        try:
            # prefer explicit model env vars if present
            model = model_name
        except NameError:
            model = None

        try:
            if model:
                enc = tiktoken.encoding_for_model(model)
            else:
                enc = tiktoken.get_encoding("cl100k_base")
        except Exception:
            # fallback encoding if model-specific lookup fails
            enc = tiktoken.get_encoding("cl100k_base")

        return len(enc.encode(text))
    except Exception as exc:
        logging.debug("tiktoken not available or failed, using heuristic token count: %s", exc)
        # Heuristic: average ~4 characters per token for English text.
        approx = max(1, int(len(text) / 4))
        return approx
    return len(text.split())
def analyze_text(text: str) -> str:
    print(f"Text for GPT-5 analysis: {text}")
    logging.info(f'Text for GPT-5 analysis: {text[:100]}')
    try:
        cfg = get_gpt5_client()
        client = cfg["client"]
        deployment = cfg["deployment"]
        model_name = cfg["model_name"]
        with open('./ai/gpt5_prompt.txt', 'r') as f:
            prompt = f.read()
    except FileNotFoundError as exc:
        logging.error("Prompt file missing: %s", exc)
        return f"Failed to load GPT-5 prompt file: {exc}"
    except Exception as exc:
        logging.error("Unexpected error loading prompt: %s", exc)
        return f"Failed to prepare GPT-5 request: {exc}"

    try:
        token_count = count_tokens(model_name, (prompt+text))
        logging.info(f'Text analysis. Prompt + text token count: {token_count}')
        response = client.chat.completions.create(
            messages=[
                {
                    "role": "system",
                    "content": prompt,
                },
                {
                    "role": "user",
                    "content": text,
                }
            ],
            max_tokens=16384,
            temperature=1.0,
            top_p=1.0,
            model=deployment
        )
    except Exception as exc:
        logging.error("GPT-5 chat completion failed: %s", exc)
        return f"GPT-5 completion failed: {exc}"

    logging.info(f'Text analysis response: {response.choices[0].message.content}')
    return str(response.choices[0].message.content)
def analyze_image(image_url: str) -> str:
    """
    Analyze image using GPT-5 by providing an image URL.
    This function expects a public/HTTPS image URL.
    """
    logging.info(f'Analyzing image URL: {image_url}')
    if not _is_remote_path(image_url):
        raise ValueError("analyze_image expects an HTTP(S) URL")

    try:
        cfg = get_gpt5_client()
        client = cfg["client"]
        deployment = cfg["deployment"]
        model_name = cfg["model_name"]
        with open('./ai/gpt5_img_prompt.txt', 'r') as f:
            prompt = f.read()
    except Exception as exc:
        logging.error("Failed to get GPT-5 client or prompt: %s", exc)
        return f"Failed to get GPT-5 client: {exc}"

    try:
        # Message format: system text + user with image_url object (image_url.url)
        response = client.chat.completions.create(
            messages=[
                {"role": "system", "content": [{"type": "text", "text": prompt}]},
                {"role": "user", "content": [{"type": "image_url", "image_url": {"url": image_url}}]}
            ],
            max_tokens=16384,
            stop=None,
            stream=False,
            model=deployment
        )
    except Exception as exc:
        logging.error("GPT-5 Image chat completion failed: %s", exc)
        return f"GPT-5 Image completion failed: {exc}"

    try:
        
        response = response.choices[0].message.content # + "\n" + "Input tokens used: " + str(response.usage.prompt_tokens) + "\n" + "Output tokens used: " + str(response.usage.completion_tokens)
        response = json.loads(response)
        cleaned_response = {k: v for k, v in response.items() if v != "None"}
        #logging.info(f'Image analysis cleaned response: {cleaned_response}')
        return json.dumps(cleaned_response, ensure_ascii=False, indent=2)
        
    except Exception:
        return str(response)
def upload_and_get_sas(account_url: str, blob_name: str, local_file: str, expiry_hours: int = 1) -> str:

    """
    GPT-5 works better with images (see flood on the floor, not just a shoe), but requires URL to the image. So needed to upload to blob and generate SAS.
    Uploads a local file to Azure Blob Storage using Managed Identity (Azure) or Storage Key (local)
    and generates a read-only SAS URL.

    Priority:
    1. Use STORAGE_ACCOUNT_KEY (if set) → useful for local dev/test.
    2. Otherwise, use Managed Identity / DefaultAzureCredential (in Azure).

    :param account_url: Storage account URL, e.g. "https://mystorageaccount.blob.core.windows.net"
    :param container: Name of the target container
    :param blob_name: Name of the blob to create
    :param local_file: Path to the local file to upload
    :param expiry_hours: SAS expiry time in hours
    :return: Read-only SAS URL to access the uploaded blob
    """

    # Check for local dev override
    account_key = os.getenv("STORAGE_ACCOUNT_KEY")
    email_attachments_container_name = os.getenv("EMAIL_ATTACHMENTS_CONTAINER")

    if not account_key:
        try:
            with open("local.settings.json", "r") as fh:
                settings = json.load(fh)["Values"]
            account_key = account_key or settings.get("STORAGE_ACCOUNT_KEY", "")
            email_attachments_container_name = email_attachments_container_name or settings.get("EMAIL_ATTACHMENTS_CONTAINER", "")
        except FileNotFoundError:
            pass

    if not email_attachments_container_name:
        raise ValueError("EMAIL_ATTACHMENTS_CONTAINER is not configured")

    if account_key:
        # Use Storage Key (local testing)
        blob_service_client = BlobServiceClient(account_url=account_url, credential=account_key)
        user_delegation_key = None  # SAS will be signed with account key
    else:
        # Use Managed Identity (Azure environment)
        credential = DefaultAzureCredential()
        blob_service_client = BlobServiceClient(account_url=account_url, credential=credential)
        # Create user delegation key for SAS
        user_delegation_key = blob_service_client.get_user_delegation_key(
            key_start_time=datetime.utcnow(),
            key_expiry_time=datetime.utcnow() + timedelta(hours=expiry_hours)
        )

    # Ensure container exists
    container_client = blob_service_client.get_container_client(email_attachments_container_name)
    try:
        container_client.create_container()
    except Exception:
        pass  # ignore if exists

    # Upload file
    blob_client = container_client.get_blob_client(blob_name)
    with open(local_file, "rb") as data:
        blob_client.upload_blob(data, overwrite=True)

    if blob_service_client.account_name:
        account_name = blob_service_client.account_name
    else:
        account_name = os.getenv("STORAGE_ACCOUNT_NAME")
    
    if not account_name:
        raise ValueError("Storage account name could not be determined")
    
    # Generate SAS token
    if account_key and blob_service_client.account_name:
        sas_token = generate_blob_sas(
            account_name=account_name,
            container_name=email_attachments_container_name,
            blob_name=blob_name,
            account_key=account_key,
            permission=BlobSasPermissions(read=True),
            expiry=datetime.utcnow() + timedelta(hours=expiry_hours)
        )
    else:
        sas_token = generate_blob_sas(
            
            account_name=account_name,
            container_name=email_attachments_container_name,
            blob_name=blob_name,
            user_delegation_key=user_delegation_key,
            permission=BlobSasPermissions(read=True),
            expiry=datetime.utcnow() + timedelta(hours=expiry_hours)
        )

    return f"{blob_client.url}?{sas_token}"



#file_path = "https://staiclaimsauto001.blob.core.windows.net/emailattachments/20251003154607_photos.docx?sp=r&st=2025-10-04T11:03:40Z&se=2025-10-11T19:18:40Z&spr=https&sv=2024-11-04&sr=b&sig=%2FvFXazh2nLE%2BAy0ovcVQTNexNcFRWaP7YznxQZZkt9M%3D"
file_path = "/emailattachments/20251003154607_photos.docx"
print(extract_file_info(file_path))



